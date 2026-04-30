import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import io

# ─────────────────────────────────────────────────────────────────────────────
# STYLE SETTINGS — edit these to change document standards
# ─────────────────────────────────────────────────────────────────────────────
STYLE = {
    "font_name":          "Times New Roman",
    "font_size":          12,
    "heading_font":       "Times New Roman",
    "heading_size":       16,
    "alignment":          "Justified",
    "line_spacing":       1.15,
    "space_before_para":  0,
    "space_after_para":   6,
    "first_line_indent":  0.0,
    "margin_top":         1.0,
    "margin_bottom":      1.0,
    "margin_left":        1.0,
    "margin_right":       1.0,
    "heading_bold":       True,
    "body_bold":          False,
    "body_italic":        False,
    "body_color":         "#000000",
    "heading_color":      "#000000",
    "include_header":     False,
    "header_text":        "",
    "include_footer":     False,
    "footer_text":        "",
    "page_numbers":       False,
}

ALIGN_MAP = {
    "Left":      WD_ALIGN_PARAGRAPH.LEFT,
    "Centered":  WD_ALIGN_PARAGRAPH.CENTER,
    "Right":     WD_ALIGN_PARAGRAPH.RIGHT,
    "Justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

HTML_ALIGN_MAP = {
    "Left": "left", "Centered": "center",
    "Right": "right", "Justified": "justify",
}

def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def build_docx(text):
    cfg = STYLE
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(cfg["margin_top"])
        section.bottom_margin = Inches(cfg["margin_bottom"])
        section.left_margin   = Inches(cfg["margin_left"])
        section.right_margin  = Inches(cfg["margin_right"])

    body_r, body_g, body_b = hex_to_rgb(cfg["body_color"])
    head_r, head_g, head_b = hex_to_rgb(cfg["heading_color"])
    alignment = ALIGN_MAP.get(cfg["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

    for line in text.splitlines():
        s = line.strip()
        if s.startswith("###"):
            p = doc.add_paragraph()
            r = p.add_run(s.lstrip("#").strip())
            r.bold = True; r.font.name = cfg["heading_font"]
            r.font.size = Pt(max(cfg["heading_size"]-4, 11))
            r.font.color.rgb = RGBColor(head_r, head_g, head_b)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(3)
        elif s.startswith("##"):
            p = doc.add_paragraph()
            r = p.add_run(s.lstrip("#").strip())
            r.bold = True; r.font.name = cfg["heading_font"]
            r.font.size = Pt(max(cfg["heading_size"]-2, 13))
            r.font.color.rgb = RGBColor(head_r, head_g, head_b)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
        elif s.startswith("#"):
            p = doc.add_paragraph()
            r = p.add_run(s.lstrip("#").strip())
            r.bold = cfg["heading_bold"]; r.font.name = cfg["heading_font"]
            r.font.size = Pt(cfg["heading_size"])
            r.font.color.rgb = RGBColor(head_r, head_g, head_b)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)
        else:
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.bold = cfg["body_bold"]; r.italic = cfg["body_italic"]
            r.font.name = cfg["font_name"]; r.font.size = Pt(cfg["font_size"])
            r.font.color.rgb = RGBColor(body_r, body_g, body_b)
            p.alignment = alignment
            pf = p.paragraph_format
            pf.space_before = Pt(cfg["space_before_para"])
            pf.space_after  = Pt(cfg["space_after_para"])
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = cfg["line_spacing"]
            if cfg["first_line_indent"] > 0:
                pf.first_line_indent = Inches(cfg["first_line_indent"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_preview(text):
    cfg = STYLE
    align = HTML_ALIGN_MAP.get(cfg["alignment"], "left")
    font  = cfg["font_name"]
    fs    = cfg["font_size"]
    hs    = cfg["heading_size"]
    ls    = cfg["line_spacing"]
    sa    = cfg["space_after_para"]
    ind   = f"{cfg['first_line_indent']}in" if cfg["first_line_indent"] > 0 else "0"
    bc    = cfg["body_color"]
    hc    = cfg["heading_color"]
    hb    = "bold" if cfg["heading_bold"] else "normal"
    bb    = "bold" if cfg["body_bold"] else "normal"
    bi    = "italic" if cfg["body_italic"] else "normal"

    body_css = (f"font-family:'{font}',serif;font-size:{fs}pt;color:{bc};"
                f"text-align:{align};line-height:{ls};margin:0 0 {sa}pt 0;"
                f"text-indent:{ind};font-weight:{bb};font-style:{bi};")

    rows = []
    for line in text.splitlines():
        s = line.strip()
        if s.startswith("###"):
            rows.append(f"<p style=\"font-family:'{font}',serif;font-size:{max(hs-4,11)}pt;"
                        f"color:{hc};font-weight:bold;margin:6pt 0 3pt;\">{s.lstrip('#').strip()}</p>")
        elif s.startswith("##"):
            rows.append(f"<p style=\"font-family:'{font}',serif;font-size:{max(hs-2,13)}pt;"
                        f"color:{hc};font-weight:bold;margin:8pt 0 4pt;\">{s.lstrip('#').strip()}</p>")
        elif s.startswith("#"):
            rows.append(f"<p style=\"font-family:'{font}',serif;font-size:{hs}pt;"
                        f"color:{hc};font-weight:{hb};margin:12pt 0 6pt;\">{s.lstrip('#').strip()}</p>")
        elif s == "":
            rows.append("<div style='height:8pt'></div>")
        else:
            rows.append(f"<p style='{body_css}'>{s}</p>")

    return f"""
    <div style="background:#fff;padding:32px 48px;border-radius:6px;
                box-shadow:0 2px 8px rgba(0,0,0,0.10);min-height:300px;
                font-family:'{font}',serif;">
        {''.join(rows) if rows else '<p style="color:#ccc;font-style:italic;">Preview will appear here...</p>'}
    </div>"""


# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="DocuStyle", page_icon="📄", layout="wide")

st.markdown("""
<style>
/* ── Global ── */
.stApp, [data-testid="stAppViewContainer"] { background: #1C1C1E; }
[data-testid="stHeader"] { background: transparent; }
section[data-testid="stSidebar"] { display: none; }

/* ── Top bar ── */
.topbar {
    display: flex;
    align-items: center;
    justify-content: space-between;
    background: #2C2C2E;
    border-radius: 10px;
    padding: 14px 24px;
    margin-bottom: 18px;
    border: 1px solid #3A3A3C;
}
.topbar-title { color: #fff; font-size: 1.3rem; font-weight: 700; letter-spacing: 0.02em; }
.topbar-sub   { color: #C4956A; font-size: 0.82rem; margin-top: 2px; }

/* ── Column cards ── */
.col-card {
    background: #2C2C2E;
    border-radius: 10px;
    border: 1px solid #3A3A3C;
    padding: 16px;
    height: 100%;
}
.col-label {
    font-size: 0.7rem;
    font-weight: 700;
    color: #C4956A;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    border-bottom: 1px solid #3A3A3C;
    padding-bottom: 8px;
    margin-bottom: 12px;
}

/* ── Textarea ── */
.stTextArea textarea {
    background: #1C1C1E !important;
    color: #F0F0F0 !important;
    border: 1px solid #3A3A3C !important;
    border-radius: 6px !important;
    font-family: 'Consolas', monospace !important;
    font-size: 0.88rem !important;
    resize: none !important;
    caret-color: #C4956A !important;
}
.stTextArea textarea:focus {
    border-color: #6B3F1F !important;
    box-shadow: 0 0 0 2px rgba(107,63,31,0.3) !important;
}

/* ── File uploader ── */
[data-testid="stFileUploader"] {
    background: #1C1C1E !important;
    border: 1px dashed #3A3A3C !important;
    border-radius: 6px !important;
    padding: 6px !important;
}
[data-testid="stFileUploader"] * { color: #888 !important; font-size: 0.8rem !important; }

/* ── Buttons ── */
.stButton > button {
    width: 100% !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    font-size: 0.88rem !important;
    padding: 9px 0 !important;
    border: none !important;
    cursor: pointer !important;
    transition: opacity 0.15s !important;
}
.stButton > button:hover { opacity: 0.85 !important; }

/* Primary - Generate */
div[data-testid="column"] .stButton:nth-of-type(1) > button {
    background: #6B3F1F !important;
    color: #fff !important;
}
/* Danger - Clear */
.clear-btn > div > button, .clear-btn button {
    background: #3A3A3C !important;
    color: #FF6B6B !important;
    border: 1px solid #FF6B6B !important;
}

/* ── Download button ── */
.stDownloadButton > button {
    width: 100% !important;
    background: #6B3F1F !important;
    color: #fff !important;
    border: none !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    font-size: 0.88rem !important;
    padding: 9px 0 !important;
}
.stDownloadButton > button:hover { opacity: 0.85 !important; }

/* ── Feedback textarea ── */
.feedback-area textarea {
    background: #1C1C1E !important;
    color: #F0F0F0 !important;
    border: 1px solid #3A3A3C !important;
    border-radius: 6px !important;
    font-size: 0.85rem !important;
    resize: none !important;
}

/* ── Counter ── */
.counter { color: #666; font-size: 0.75rem; text-align: right; margin-top: -6px; }

/* ── Empty preview ── */
.preview-empty {
    display: flex; align-items: center; justify-content: center;
    height: 300px; color: #555; font-size: 0.9rem; font-style: italic;
}

/* ── Scrollable preview ── */
.preview-scroll {
    max-height: 560px;
    overflow-y: auto;
    border-radius: 6px;
}

/* ── Labels ── */
label, .stTextArea label { color: #888 !important; font-size: 0.8rem !important; }

/* hide streamlit chrome */
#MainMenu, footer, header { visibility: hidden; }
div[data-testid="stDecoration"] { display: none; }
</style>
""", unsafe_allow_html=True)

# ── Init session state ─────────────────────────────────────────────────────────
if "text_input" not in st.session_state:
    st.session_state.text_input = ""
if "preview_html" not in st.session_state:
    st.session_state.preview_html = ""
if "docx_bytes" not in st.session_state:
    st.session_state.docx_bytes = None
if "clear_trigger" not in st.session_state:
    st.session_state.clear_trigger = 0

# ── Top bar ───────────────────────────────────────────────────────────────────
st.markdown("""
<div class="topbar">
    <div>
        <div class="topbar-title">📄 DocuStyle</div>
        <div class="topbar-sub">Internal Document Formatting Tool</div>
    </div>
</div>
""", unsafe_allow_html=True)

# ── Three columns ─────────────────────────────────────────────────────────────
left, mid, right = st.columns([1.1, 1.4, 0.9])

# ─── LEFT: Input ──────────────────────────────────────────────────────────────
with left:
    st.markdown('<div class="col-card">', unsafe_allow_html=True)
    st.markdown('<div class="col-label">Input</div>', unsafe_allow_html=True)

    uploaded = st.file_uploader("Upload .txt file", type=["txt"], label_visibility="collapsed")
    if uploaded:
        st.session_state.text_input = uploaded.read().decode("utf-8", errors="replace")

    # Generate button at top
    if st.button("⚡ Generate Preview", key="generate"):
        if st.session_state.text_input.strip():
            st.session_state.preview_html = build_preview(st.session_state.text_input)
            st.session_state.docx_bytes   = build_docx(st.session_state.text_input)
        else:
            st.warning("Please enter some text first.")

    text_val = st.text_area(
        "Paste text:",
        value=st.session_state.text_input,
        height=460,
        key=f"textarea_{st.session_state.clear_trigger}",
        label_visibility="collapsed",
        placeholder="Paste your text here..."
    )
    st.session_state.text_input = text_val

    char_count = len(text_val)
    line_count = text_val.count("\n") + 1 if text_val.strip() else 0
    st.markdown(f'<div class="counter">{char_count:,} chars · {line_count:,} lines</div>',
                unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

# ─── MIDDLE: Preview ──────────────────────────────────────────────────────────
with mid:
    st.markdown('<div class="col-card">', unsafe_allow_html=True)
    st.markdown('<div class="col-label">Preview</div>', unsafe_allow_html=True)

    if st.session_state.preview_html:
        st.components.v1.html(
            f'<div class="preview-scroll">{st.session_state.preview_html}</div>',
            height=560,
            scrolling=True
        )
    else:
        st.markdown("""
        <div class="preview-empty">
            Paste your text and click Generate Preview
        </div>
        """, unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

# ─── RIGHT: Actions ───────────────────────────────────────────────────────────
with right:
    st.markdown('<div class="col-card">', unsafe_allow_html=True)
    st.markdown('<div class="col-label">Actions</div>', unsafe_allow_html=True)

    # Download
    if st.session_state.docx_bytes:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        st.download_button(
            label="⬇️ Download .docx",
            data=st.session_state.docx_bytes,
            file_name=f"document_{ts}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.markdown("""
        <div style="background:#1C1C1E;border:1px solid #3A3A3C;border-radius:6px;
                    padding:9px;text-align:center;color:#555;font-size:0.85rem;">
            ⬇️ Download .docx
        </div>""", unsafe_allow_html=True)

    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

    # Clear
    st.markdown('<div class="clear-btn">', unsafe_allow_html=True)
    if st.button("🗑️ Clear Text", key="clear"):
        st.session_state.text_input  = ""
        st.session_state.preview_html = ""
        st.session_state.docx_bytes  = None
        st.session_state.clear_trigger += 1
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)

    # Feedback
    st.markdown('<div class="col-label">Feedback</div>', unsafe_allow_html=True)
    st.markdown('<div class="feedback-area">', unsafe_allow_html=True)
    feedback = st.text_area("Share your thoughts on the output:",
                            height=160,
                            placeholder="How did the formatting look? Any issues?",
                            label_visibility="collapsed",
                            key="feedback_box")
    st.markdown('</div>', unsafe_allow_html=True)

    if st.button("📨 Submit Feedback", key="submit_feedback"):
        if feedback.strip():
            # PLACEHOLDER: replace this with your shared file path when ready
            # e.g. with open(r"\\network\share\feedback.txt", "a") as f:
            #          f.write(f"{datetime.datetime.now()} — {feedback}\n")
            st.success("Thanks for your feedback!")
        else:
            st.warning("Please write something before submitting.")

    st.markdown('</div>', unsafe_allow_html=True)
