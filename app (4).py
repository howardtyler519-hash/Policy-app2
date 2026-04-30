import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import io

# ─────────────────────────────────────────────────────────────────────────────
# STYLE SETTINGS — edit these to change the document standards
# ─────────────────────────────────────────────────────────────────────────────

STYLE = {
    "font_name":          "Times New Roman",
    "font_size":          12,
    "heading_font":       "Times New Roman",
    "heading_size":       16,
    "alignment":          "Justified",
    "line_spacing":       1.15,
    "space_before_para":  0,
    "space_after_para":   6,
    "first_line_indent":  0.0,
    "margin_top":         1.0,
    "margin_bottom":      1.0,
    "margin_left":        1.0,
    "margin_right":       1.0,
    "heading_bold":       True,
    "body_bold":          False,
    "body_italic":        False,
    "body_color":         "#000000",
    "heading_color":      "#000000",
    "include_header":     False,
    "header_text":        "",
    "include_footer":     False,
    "footer_text":        "",
    "page_numbers":       False,
}

# ─────────────────────────────────────────────────────────────────────────────

ALIGN_MAP = {
    "Left":      WD_ALIGN_PARAGRAPH.LEFT,
    "Centered":  WD_ALIGN_PARAGRAPH.CENTER,
    "Right":     WD_ALIGN_PARAGRAPH.RIGHT,
    "Justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

HTML_ALIGN_MAP = {
    "Left":      "left",
    "Centered":  "center",
    "Right":     "right",
    "Justified": "justify",
}

def hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def build_docx(text: str) -> bytes:
    cfg = STYLE
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(cfg["margin_top"])
        section.bottom_margin = Inches(cfg["margin_bottom"])
        section.left_margin   = Inches(cfg["margin_left"])
        section.right_margin  = Inches(cfg["margin_right"])

    if cfg["include_header"] and cfg["header_text"].strip():
        for section in doc.sections:
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.clear()
            run = p.add_run(cfg["header_text"])
            run.font.name = cfg["font_name"]
            run.font.size = Pt(cfg["font_size"] - 2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if cfg["include_footer"] or cfg["page_numbers"]:
        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cfg["include_footer"] and cfg["footer_text"].strip():
                run = p.add_run(cfg["footer_text"])
                run.font.name = cfg["font_name"]
                run.font.size = Pt(cfg["font_size"] - 2)
            if cfg["page_numbers"]:
                run = p.add_run()
                for tag, val in [("w:fldChar", "begin"), ("w:instrText", "PAGE"), ("w:fldChar", "separate"), ("w:fldChar", "end")]:
                    el = OxmlElement(tag)
                    if tag == "w:instrText":
                        el.text = val
                    else:
                        el.set(qn("w:fldCharType"), val)
                    run._r.append(el)

    body_r, body_g, body_b = hex_to_rgb(cfg["body_color"])
    head_r, head_g, head_b = hex_to_rgb(cfg["heading_color"])
    alignment = ALIGN_MAP.get(cfg["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("###"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.name = cfg["heading_font"]
            run.font.size = Pt(max(cfg["heading_size"] - 4, 11))
            run.font.color.rgb = RGBColor(head_r, head_g, head_b)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(3)
        elif stripped.startswith("##"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.name = cfg["heading_font"]
            run.font.size = Pt(max(cfg["heading_size"] - 2, 13))
            run.font.color.rgb = RGBColor(head_r, head_g, head_b)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
        elif stripped.startswith("#"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = cfg["heading_bold"]
            run.font.name = cfg["heading_font"]
            run.font.size = Pt(cfg["heading_size"])
            run.font.color.rgb = RGBColor(head_r, head_g, head_b)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)
        else:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold   = cfg["body_bold"]
            run.italic = cfg["body_italic"]
            run.font.name  = cfg["font_name"]
            run.font.size  = Pt(cfg["font_size"])
            run.font.color.rgb = RGBColor(body_r, body_g, body_b)
            p.alignment = alignment
            pf = p.paragraph_format
            pf.space_before = Pt(cfg["space_before_para"])
            pf.space_after  = Pt(cfg["space_after_para"])
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = cfg["line_spacing"]
            if cfg["first_line_indent"] > 0:
                pf.first_line_indent = Inches(cfg["first_line_indent"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_preview(text: str) -> str:
    cfg = STYLE
    align = HTML_ALIGN_MAP.get(cfg["alignment"], "left")
    body_color = cfg["body_color"]
    head_color = cfg["heading_color"]
    font = cfg["font_name"]
    font_size = cfg["font_size"]
    heading_size = cfg["heading_size"]
    line_spacing = cfg["line_spacing"]
    space_after = cfg["space_after_para"]
    indent = f"{cfg['first_line_indent']}in" if cfg["first_line_indent"] > 0 else "0"

    body_style = (
        f"font-family:'{font}', serif;"
        f"font-size:{font_size}pt;"
        f"color:{body_color};"
        f"text-align:{align};"
        f"line-height:{line_spacing};"
        f"margin-bottom:{space_after}pt;"
        f"margin-top:0;"
        f"text-indent:{indent};"
        f"{'font-weight:bold;' if cfg['body_bold'] else ''}"
        f"{'font-style:italic;' if cfg['body_italic'] else ''}"
    )

    html_lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("###"):
            content = stripped.lstrip("#").strip()
            html_lines.append(
                f"<p style=\"font-family:'{font}',serif;font-size:{max(heading_size-4,11)}pt;"
                f"color:{head_color};font-weight:bold;margin-bottom:3pt;margin-top:6pt;\">"
                f"{content}</p>"
            )
        elif stripped.startswith("##"):
            content = stripped.lstrip("#").strip()
            html_lines.append(
                f"<p style=\"font-family:'{font}',serif;font-size:{max(heading_size-2,13)}pt;"
                f"color:{head_color};font-weight:bold;margin-bottom:4pt;margin-top:8pt;\">"
                f"{content}</p>"
            )
        elif stripped.startswith("#"):
            content = stripped.lstrip("#").strip()
            html_lines.append(
                f"<p style=\"font-family:'{font}',serif;font-size:{heading_size}pt;"
                f"color:{head_color};font-weight:{'bold' if cfg['heading_bold'] else 'normal'};"
                f"margin-bottom:6pt;margin-top:12pt;\">"
                f"{content}</p>"
            )
        elif stripped == "":
            html_lines.append("<br>")
        else:
            html_lines.append(f"<p style=\"{body_style}\">{stripped}</p>")

    margin = f"{cfg['margin_left']}in"
    preview_html = f"""
    <div style="
        background:white;
        padding: 60px {margin};
        max-width: 850px;
        margin: 0 auto;
        box-shadow: 0 2px 12px rgba(0,0,0,0.12);
        border-radius: 4px;
        min-height: 400px;
    ">
        {''.join(html_lines)}
    </div>
    """
    return preview_html


# ─────────────────────────────────────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="DocuStyle",
    page_icon="📄",
    layout="centered"
)

st.markdown("""
<style>
    .stApp { background-color: #F5F0EB; }
    [data-testid="stAppViewContainer"] { background-color: #F5F0EB; }

    .app-header {
        background-color: #6B3F1F;
        padding: 20px 28px;
        border-radius: 8px;
        margin-bottom: 24px;
    }
    .app-header h1 {
        color: #FFFFFF !important;
        margin: 0 !important;
        font-size: 1.8rem !important;
        font-weight: 700 !important;
    }
    .app-header p {
        color: #C4956A !important;
        margin: 4px 0 0 0 !important;
        font-size: 0.9rem !important;
    }

    .section-box {
        background-color: #FFFFFF;
        border-radius: 8px;
        padding: 24px;
        margin-bottom: 20px;
        box-shadow: 0 1px 4px rgba(0,0,0,0.07);
    }

    .section-label {
        font-size: 0.75rem;
        font-weight: 700;
        color: #7A4F2D;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        margin-bottom: 12px;
        border-bottom: 1px solid #C4956A;
        padding-bottom: 6px;
    }

    .stTextArea textarea {
        border: 1.5px solid #C4956A !important;
        border-radius: 6px !important;
        font-family: 'Consolas', monospace !important;
        font-size: 0.9rem !important;
        background-color: #FAFAFA !important;
    }
    .stTextArea textarea:focus {
        border-color: #6B3F1F !important;
        box-shadow: 0 0 0 2px rgba(107,63,31,0.15) !important;
    }

    .stDownloadButton button {
        background-color: #6B3F1F !important;
        color: #FFFFFF !important;
        border: none !important;
        border-radius: 6px !important;
        padding: 12px 28px !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        width: 100% !important;
    }
    .stDownloadButton button:hover {
        background-color: #9B6B47 !important;
    }

    .stButton button {
        background-color: #FFFFFF !important;
        color: #6B3F1F !important;
        border: 1.5px solid #6B3F1F !important;
        border-radius: 6px !important;
        padding: 10px 28px !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        width: 100% !important;
    }
    .stButton button:hover {
        background-color: #FAF7F4 !important;
    }

    .tip-box {
        background-color: #FAF7F4;
        border-left: 4px solid #C4956A;
        padding: 10px 16px;
        border-radius: 4px;
        font-size: 0.85rem;
        color: #555;
        margin-bottom: 4px;
    }

    .counter {
        text-align: right;
        color: #888;
        font-size: 0.8rem;
        margin-top: -8px;
        margin-bottom: 8px;
    }

    .feedback-note {
        background-color: #FAF7F4;
        border: 1px dashed #C4956A;
        border-radius: 6px;
        padding: 16px;
        color: #888;
        font-size: 0.85rem;
        text-align: center;
    }

    #MainMenu, footer, header { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ── Header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="app-header">
    <h1>📄 DocuStyle</h1>
    <p>Paste or upload your text, preview the formatting, then download your Word document.</p>
</div>
""", unsafe_allow_html=True)

# ── Input Section ─────────────────────────────────────────────────────────────
st.markdown('<div class="section-box">', unsafe_allow_html=True)
st.markdown('<div class="section-label">Input</div>', unsafe_allow_html=True)

st.markdown("""
<div class="tip-box">
    <strong>Heading tip:</strong> Start a line with <code>#</code> for a main heading, 
    <code>##</code> for a subheading, <code>###</code> for a smaller subheading.
</div>
""", unsafe_allow_html=True)

uploaded = st.file_uploader("Upload a .txt file (optional)", type=["txt"], label_visibility="collapsed")
default_text = ""
if uploaded:
    default_text = uploaded.read().decode("utf-8", errors="replace")

text_input = st.text_area(
    "Paste your text here:",
    value=default_text,
    height=320,
    placeholder="Start typing or paste your text here...\n\n# Use # for headings\n## Use ## for subheadings",
    label_visibility="collapsed"
)

char_count = len(text_input)
line_count = text_input.count("\n") + 1 if text_input.strip() else 0
st.markdown(f'<div class="counter">{char_count:,} characters &nbsp;|&nbsp; {line_count:,} lines</div>',
            unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)

# ── Preview & Download ────────────────────────────────────────────────────────
if text_input.strip():
    # Preview
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">Document Preview</div>', unsafe_allow_html=True)
    preview_html = build_preview(text_input)
    st.components.v1.html(preview_html, height=500, scrolling=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # Download
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">Download</div>', unsafe_allow_html=True)
    docx_bytes = build_docx(text_input)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    st.download_button(
        label="⬇️  Download Word Document",
        data=docx_bytes,
        file_name=f"document_{ts}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # Feedback
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">Feedback</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class="feedback-note">
        💬 Feedback coming soon — this section will let you share thoughts on the output.
    </div>
    """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

else:
    st.markdown("""
    <div style="text-align:center; color:#AAA; padding: 40px 0; font-size:0.95rem;">
        Enter some text above to see a preview and download your document.
    </div>
    """, unsafe_allow_html=True)

st.markdown('<p style="text-align:center; color:#AAA; font-size:0.8rem; margin-top:8px;">DocuStyle — Internal Document Tool</p>',
            unsafe_allow_html=True)
