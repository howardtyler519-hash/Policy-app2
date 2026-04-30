import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import io

# ─────────────────────────────────────────────────────────────────────────────
# STYLE SETTINGS — edit these to change the document standards
# ─────────────────────────────────────────────────────────────────────────────

STYLE = {
    "font_name":          "Times New Roman",   # Body font
    "font_size":          12,                  # Body font size (pt)
    "heading_font":       "Times New Roman",   # Heading font
    "heading_size":       16,                  # H1 size (pt). H2 = -2, H3 = -4
    "alignment":          "Justified",         # Left | Centered | Right | Justified
    "line_spacing":       1.15,                # 1.0 = single, 1.5, 2.0 = double
    "space_before_para":  0,                   # Space before paragraph (pt)
    "space_after_para":   6,                   # Space after paragraph (pt)
    "first_line_indent":  0.0,                 # First line indent (inches). 0 = none
    "margin_top":         1.0,                 # Page margin top (inches)
    "margin_bottom":      1.0,                 # Page margin bottom (inches)
    "margin_left":        1.0,                 # Page margin left (inches)
    "margin_right":       1.0,                 # Page margin right (inches)
    "heading_bold":       True,                # Headings bold?
    "body_bold":          False,               # Body text bold?
    "body_italic":        False,               # Body text italic?
    "body_color":         "#000000",           # Body text color (hex)
    "heading_color":      "#000000",           # Heading color (hex)
    "include_header":     False,               # Show header on each page?
    "header_text":        "",                  # Header text
    "include_footer":     False,               # Show footer on each page?
    "footer_text":        "",                  # Footer text
    "page_numbers":       False,               # Show page numbers in footer?
}

# ─────────────────────────────────────────────────────────────────────────────

ALIGN_MAP = {
    "Left":      WD_ALIGN_PARAGRAPH.LEFT,
    "Centered":  WD_ALIGN_PARAGRAPH.CENTER,
    "Right":     WD_ALIGN_PARAGRAPH.RIGHT,
    "Justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

def hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def build_docx(text: str) -> bytes:
    cfg = STYLE
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(cfg["margin_top"])
        section.bottom_margin = Inches(cfg["margin_bottom"])
        section.left_margin   = Inches(cfg["margin_left"])
        section.right_margin  = Inches(cfg["margin_right"])

    # Header
    if cfg["include_header"] and cfg["header_text"].strip():
        for section in doc.sections:
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.clear()
            run = p.add_run(cfg["header_text"])
            run.font.name = cfg["font_name"]
            run.font.size = Pt(cfg["font_size"] - 2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer
    if cfg["include_footer"] or cfg["page_numbers"]:
        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cfg["include_footer"] and cfg["footer_text"].strip():
                run = p.add_run(cfg["footer_text"])
                run.font.name = cfg["font_name"]
                run.font.size = Pt(cfg["font_size"] - 2)
                if cfg["page_numbers"]:
                    p.add_run("  |  Page ")
            if cfg["page_numbers"]:
                run = p.add_run()
                for tag, val in [("w:fldChar", "begin"), ("w:instrText", "PAGE"), ("w:fldChar", "separate"), ("w:fldChar", "end")]:
                    el = OxmlElement(tag)
                    if tag == "w:instrText":
                        el.text = val
                    else:
                        el.set(qn("w:fldCharType"), val)
                    run._r.append(el)

    body_r, body_g, body_b = hex_to_rgb(cfg["body_color"])
    head_r, head_g, head_b = hex_to_rgb(cfg["heading_color"])
    alignment = ALIGN_MAP.get(cfg["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

    for line in text.splitlines():
        stripped = line.strip()

        if stripped.startswith("###"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.name = cfg["heading_font"]
            run.font.size = Pt(max(cfg["heading_size"] - 4, 11))
            run.font.color.rgb = RGBColor(head_r, head_g, head_b)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(3)

        elif stripped.startswith("##"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.name = cfg["heading_font"]
            run.font.size = Pt(max(cfg["heading_size"] - 2, 13))
            run.font.color.rgb = RGBColor(head_r, head_g, head_b)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)

        elif stripped.startswith("#"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = cfg["heading_bold"]
            run.font.name = cfg["heading_font"]
            run.font.size = Pt(cfg["heading_size"])
            run.font.color.rgb = RGBColor(head_r, head_g, head_b)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)

        else:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold   = cfg["body_bold"]
            run.italic = cfg["body_italic"]
            run.font.name  = cfg["font_name"]
            run.font.size  = Pt(cfg["font_size"])
            run.font.color.rgb = RGBColor(body_r, body_g, body_b)
            p.alignment = alignment
            pf = p.paragraph_format
            pf.space_before = Pt(cfg["space_before_para"])
            pf.space_after  = Pt(cfg["space_after_para"])
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = cfg["line_spacing"]
            if cfg["first_line_indent"] > 0:
                pf.first_line_indent = Inches(cfg["first_line_indent"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="DocuStyle",
    page_icon="📄",
    layout="centered"
)

st.markdown("""
<style>
    /* Background */
    .stApp { background-color: #FFFFFF; }
    [data-testid="stAppViewContainer"] { background-color: #FFFFFF; }

    /* Header bar */
    .app-header {
        background-color: #6B3F1F;
        padding: 18px 28px;
        border-radius: 8px;
        margin-bottom: 28px;
    }
    .app-header h1 {
        color: #FFFFFF !important;
        margin: 0 !important;
        font-size: 1.8rem !important;
        font-weight: 700 !important;
    }
    .app-header p {
        color: #C4956A !important;
        margin: 4px 0 0 0 !important;
        font-size: 0.9rem !important;
    }

    /* Text area */
    .stTextArea textarea {
        border: 1.5px solid #C4956A !important;
        border-radius: 6px !important;
        font-family: 'Consolas', monospace !important;
        font-size: 0.9rem !important;
        background-color: #FAFAFA !important;
    }
    .stTextArea textarea:focus {
        border-color: #6B3F1F !important;
        box-shadow: 0 0 0 2px rgba(107,63,31,0.15) !important;
    }

    /* Button */
    .stDownloadButton button, .stButton button {
        background-color: #6B3F1F !important;
        color: #FFFFFF !important;
        border: none !important;
        border-radius: 6px !important;
        padding: 10px 28px !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        width: 100% !important;
        cursor: pointer !important;
    }
    .stDownloadButton button:hover, .stButton button:hover {
        background-color: #9B6B47 !important;
    }

    /* Tip box */
    .tip-box {
        background-color: #FAF7F4;
        border-left: 4px solid #C4956A;
        padding: 10px 16px;
        border-radius: 4px;
        font-size: 0.85rem;
        color: #555;
        margin-bottom: 16px;
    }

    /* File uploader */
    [data-testid="stFileUploader"] {
        border: 1.5px dashed #C4956A !important;
        border-radius: 6px !important;
        background-color: #FAF7F4 !important;
    }

    /* Hide streamlit branding */
    #MainMenu, footer, header { visibility: hidden; }

    /* Counter */
    .counter {
        text-align: right;
        color: #888;
        font-size: 0.8rem;
        margin-top: -12px;
        margin-bottom: 12px;
    }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown("""
<div class="app-header">
    <h1>📄 DocuStyle</h1>
    <p>Paste or upload your text and download a formatted Word document instantly.</p>
</div>
""", unsafe_allow_html=True)

# Tip
st.markdown("""
<div class="tip-box">
    <strong>Heading tip:</strong> Start a line with <code>#</code> for a main heading, 
    <code>##</code> for a subheading, <code>###</code> for a smaller subheading.
</div>
""", unsafe_allow_html=True)

# File upload
uploaded = st.file_uploader("Upload a .txt file (optional)", type=["txt"])

# Text area
default_text = ""
if uploaded:
    default_text = uploaded.read().decode("utf-8", errors="replace")

text_input = st.text_area(
    "Or paste your text here:",
    value=default_text,
    height=380,
    placeholder="Start typing or paste your text here...\n\n# Use # for headings\n## Use ## for subheadings"
)

# Counter
char_count = len(text_input)
line_count = text_input.count("\n") + 1 if text_input.strip() else 0
st.markdown(f'<div class="counter">{char_count:,} characters &nbsp;|&nbsp; {line_count:,} lines</div>',
            unsafe_allow_html=True)

# Convert
if text_input.strip():
    docx_bytes = build_docx(text_input)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    st.download_button(
        label="⬇️  Download Word Document",
        data=docx_bytes,
        file_name=f"document_{ts}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.button("⬇️  Download Word Document", disabled=True)

st.markdown("---")
st.markdown('<p style="text-align:center; color:#AAA; font-size:0.8rem;">DocuStyle — Internal Document Tool</p>',
            unsafe_allow_html=True)
